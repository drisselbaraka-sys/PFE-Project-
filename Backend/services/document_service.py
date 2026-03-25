import io
from typing import List
import re

MAX_TEXT_PER_FILE_CHARS = 120_000
MAX_TOTAL_TEXT_CHARS    = 300_000
MAX_DIGEST_CHARS        = 8_000
MAX_DIGEST_LINES        = 40

NON_TECHNICAL_METADATA_PATTERNS = [
    r"\bauteur\b",
    r"\bautrice\b",
    r"\bauthor\b",
    r"\bpubli[ée]\b",
    r"\bpublication\b",
    r"\bdate\b",
    r"\bmis\s+[àa]\s+jour\b",
    r"\bversion\b",
    r"\bcopyright\b",
    r"\btous droits r[ée]serv[ée]s\b",
    r"\blicence\b",
    r"\bemail\b",
    r"\bt[ée]l[ée]phone\b",
    r"\bcontact\b",
    r"\buniversit[ée]\b",
    r"\bfacult[ée]\b",
]


def _looks_like_non_technical_metadata(text: str) -> bool:
    compact = re.sub(r"\s+", " ", (text or "")).strip().lower()
    if not compact:
        return False

    if len(compact) <= 120 and any(re.search(pattern, compact) for pattern in NON_TECHNICAL_METADATA_PATTERNS):
        return True

    if re.search(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", compact):
        return True

    return False


def _normalize_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\x00", " ")
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def _truncate_text(text: str, max_chars: int) -> str:
    if not text or len(text) <= max_chars:
        return text or ""
    return text[:max_chars]


def _sentence_score(sentence: str) -> int:
    if _looks_like_non_technical_metadata(sentence):
        return -10

    score = 0
    if re.search(r"\d", sentence):
        score += 3
    if any(token in sentence for token in (":", ";", "->", "%", "=", "(", ")")):
        score += 2
    words = sentence.split()
    if 5 <= len(words) <= 28:
        score += 2
    if sentence[:1].isupper():
        score += 1
    return score


def _extract_digest(text: str) -> str:
    normalized = _normalize_text(text)
    if not normalized:
        return ""

    paragraphs = [paragraph.strip() for paragraph in normalized.split("\n") if paragraph.strip()]
    if not paragraphs:
        return ""

    headings = []
    bullets = []
    candidates = []

    for paragraph in paragraphs:
        compact = re.sub(r"\s+", " ", paragraph).strip()
        if len(compact) < 8:
            continue
        if _looks_like_non_technical_metadata(compact):
            continue

        if len(compact) <= 90 and not compact.endswith('.'):
            headings.append(compact)
            continue

        if compact.startswith(("-", "*", "•")):
            bullets.append(compact.lstrip("-*• "))
            continue

        sentences = re.split(r"(?<=[\.!?])\s+", compact)
        for sentence in sentences:
            sentence = sentence.strip()
            if 25 <= len(sentence) <= 260:
                candidates.append(sentence)

    scored = sorted(candidates, key=lambda item: (-_sentence_score(item), len(item)))

    selected = []
    seen = set()
    for group in (headings[:8], bullets[:12], scored[:24]):
        for item in group:
            key = item.lower()
            if key in seen:
                continue
            seen.add(key)
            selected.append(item)
            if len(selected) >= MAX_DIGEST_LINES:
                break
        if len(selected) >= MAX_DIGEST_LINES:
            break

    digest = "\n".join(f"- {item}" for item in selected)
    return _truncate_text(digest, MAX_DIGEST_CHARS)


class DocumentService:

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        text = ""
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                try:
                    page_text = page.extract_text() or ""
                except Exception:
                    page_text = ""
                if page_text:
                    text += page_text + "\n"
                if len(text) >= MAX_TEXT_PER_FILE_CHARS:
                    break
            print(f"[DocService] pypdf: {len(text)} chars extraits")
        except Exception as e:
            print(f"[DocService] pypdf failed: {e}")

        return _truncate_text(_normalize_text(text), MAX_TEXT_PER_FILE_CHARS)

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        text = ""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"
                if len(text) >= MAX_TEXT_PER_FILE_CHARS:
                    break
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                if len(text) >= MAX_TEXT_PER_FILE_CHARS:
                    break
        except Exception as e:
            print(f"[DocService] DOCX error: {e}")
        return _truncate_text(_normalize_text(text), MAX_TEXT_PER_FILE_CHARS)

    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> str:
        decoded = ""
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                decoded = file_bytes.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        return _truncate_text(_normalize_text(decoded), MAX_TEXT_PER_FILE_CHARS)

    @classmethod
    def process_files(cls, files: List[tuple]) -> str:
        digests = []

        for filename, content in files:
            fname_lower = filename.lower()
            extracted   = ""

            if fname_lower.endswith(".pdf"):
                extracted = cls.extract_text_from_pdf(content)
            elif fname_lower.endswith(".docx"):
                extracted = cls.extract_text_from_docx(content)
            elif fname_lower.endswith((".txt", ".md")):
                extracted = cls.extract_text_from_txt(content)
            else:
                print(f"[DocService] Format non supporté: {filename}")
                continue

            print(f"[DocService] '{filename}': {len(extracted)} chars extraits")

            if extracted:
                digest = _extract_digest(extracted)
                if digest:
                    digests.append(f"[DOCUMENT_SUMMARY: {filename}]\n{digest}")
                else:
                    digests.append(f"[DOCUMENT_SUMMARY: {filename}]\n{_truncate_text(extracted, 2000)}")

            if sum(len(part) for part in digests) >= MAX_TOTAL_TEXT_CHARS:
                break

        combined_text = "\n\n".join(digests)
        return _normalize_text(_truncate_text(combined_text, MAX_TOTAL_TEXT_CHARS))